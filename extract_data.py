import os
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_core.documents import Document
# -------------------------
# CONFIG
# -------------------------
BASE_DATA_DIR = r"D:\UniAgent\Data"
CHUNKS_BASE_DIR = r"D:\UniAgent\Chunks"
VECTOR_DB_BASE_DIR = r"D:\UniAgent\VectorDBs"

# Adjust these according to your LLM/token needs
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 100

# ------------------------
# HELPER FUNCTIONS
# -------------------------

import os
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

# Define chunk constants
CHUNK_SIZE = 500
CHUNK_OVERLAP = 50

def extract_chunks_from_docx(file_path, university_name):
    """
    Extracts Word document content into chunks:
    - Each heading + its content (including tables) becomes one chunk.
    - Long sections are split by RecursiveCharacterTextSplitter.
    - Returns a list of LangChain Document objects with metadata.
    """
    doc = DocxDocument(file_path)
    chunks = []

    current_heading = None
    current_content = []

    # Iterate through all elements in the document (paragraphs + tables)
    for block in doc.element.body:
        if block.tag.endswith("p"):  # Paragraph
            para = block
            text = para.text.strip() if hasattr(para, "text") else ""
            if not text:
                continue

            # Detect heading
            para_obj = next((p for p in doc.paragraphs if p._p == para), None)
            if para_obj and para_obj.style.name.startswith("Heading"):
                # Save previous heading + content
                if current_heading:
                    combined_text = current_heading + "\n" + "\n".join(current_content)
                    text_splitter = RecursiveCharacterTextSplitter(
                        chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP
                    )
                    for split_text in text_splitter.split_text(combined_text):
                        chunks.append(
                            Document(
                                page_content=split_text,
                                metadata={
                                    "university": university_name,
                                    "source_file": os.path.basename(file_path),
                                    "type": "section",
                                    "heading": current_heading,
                                },
                            )
                        )

                # Start new section
                current_heading = text
                current_content = []

            else:
                current_content.append(text)

        elif block.tag.endswith("tbl"):  # Table
            # Find corresponding docx table object
            table_obj = next((t for t in doc.tables if t._tbl == block), None)
            if table_obj:
                table_text = []
                for row in table_obj.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_text.append(" | ".join(row_data))
                # Convert table to readable markdown-like text
                current_content.append("\n".join(table_text))

    # Handle last section
    if current_heading and current_content:
        combined_text = current_heading + "\n" + "\n".join(current_content)
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP
        )
        for split_text in text_splitter.split_text(combined_text):
            chunks.append(
                Document(
                    page_content=split_text,
                    metadata={
                        "university": university_name,
                        "source_file": os.path.basename(file_path),
                        "type": "section",
                        "heading": current_heading,
                    },
                )
            )

    return chunks


def create_chunks_and_vector_db(university_name, docs_folder):
    """
    Extracts chunks from all docs in a folder, splits large chunks, and creates a FAISS vector DB.
    Stores chunks and DB in separate folders per university.
    """
    # Step 1: Extract chunks from all documents
    all_chunks = []
    for file in os.listdir(docs_folder):
        if file.endswith(".docx") and not file.startswith("~$"):
            file_path = os.path.join(docs_folder, file)
            print(f"[{university_name}] Processing {file_path}")
            chunks = extract_chunks_from_docx(file_path, university_name)
            all_chunks.extend(chunks)


    # Step 2: Further split large chunks
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP
    )

    final_docs = []
    for doc in all_chunks:
        splits = text_splitter.split_text(doc.page_content)
        for s in splits:
            final_docs.append(Document(page_content=s, metadata=doc.metadata))

    print(f"[{university_name}] Total chunks for embeddings: {len(final_docs)}")

    # Step 3: Save chunks to folder (optional)
    chunks_folder = os.path.join(CHUNKS_BASE_DIR, f"{university_name}_chunks")
    os.makedirs(chunks_folder, exist_ok=True)
    for i, doc in enumerate(final_docs):
        chunk_file = os.path.join(chunks_folder, f"chunk_{i}.txt")
        with open(chunk_file, "w", encoding="utf-8") as f:
            f.write(doc.page_content)

    # Step 4: Create vector DB using huggingface embeddings
    embeddings = HuggingFaceEmbeddings(
    model_name="sentence-transformers/all-MiniLM-L6-v2"
)
    vector_db = FAISS.from_documents(final_docs, embeddings)

    # Step 5: Save vector DB
    vector_db_path = os.path.join(VECTOR_DB_BASE_DIR, f"{university_name}_faiss")
    os.makedirs(vector_db_path, exist_ok=True)
    vector_db.save_local(vector_db_path)

    print(f"[{university_name}] Vector DB saved at {vector_db_path}")

# -------------------------
# MAIN
# -------------------------
if __name__ == "__main__":
    os.makedirs(CHUNKS_BASE_DIR, exist_ok=True)
    os.makedirs(VECTOR_DB_BASE_DIR, exist_ok=True)

    # Loop through each university folder
    for university_folder in os.listdir(BASE_DATA_DIR):
        uni_path = os.path.join(BASE_DATA_DIR, university_folder)
        if os.path.isdir(uni_path):
            create_chunks_and_vector_db(university_folder, uni_path)

    print("All universities processed successfully!")
